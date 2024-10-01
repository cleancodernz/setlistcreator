import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import yaml
import argparse
import logging

# Function to convert song time (in "MM:SS" format) to total seconds
def time_to_seconds(time_str):
    minutes, seconds = map(int, time_str.split(':'))
    return minutes * 60 + seconds

# Function to convert total seconds back to "MM:SS" format
def seconds_to_time(total_seconds):
    minutes = total_seconds // 60
    seconds = total_seconds % 60
    return f"{minutes}:{seconds:02d}"

# Function to add a song to the document with a dynamic font size
def add_song_to_document(doc, song_number, song, length, font_size):
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run(f"{song_number} : {song} - {length}")
    run.font.size = Pt(font_size)

def add_set_number(doc, set):
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run(f"Set - {set}")
    run.font.size = Pt(16)

# Add the gig details to the top of the document
def add_gig_details(doc, band_name, gig_venue, gig_date):
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run(f"{band_name} - {gig_venue}\n{gig_date}")
    run.font.size = Pt(16)
    run.bold = True

# Function to create the setlist document
def create_setlist(csv_file, config_file=None):
    # Read the CSV file containing songs and lengths
    df = pd.read_csv(csv_file)

    # Initialize a new Word document
    doc = Document()

    # If a config file is provided, read gig details from it
    if config_file:
        with open(config_file, 'r') as file:
            config = yaml.safe_load(file)
        gig_date = config.get('gig_date', 'Date not provided')
        gig_venue = config.get('gig_venue', 'Venue not provided')
        band_name = config.get('band_name', 'Band not provided')
    else:
        gig_date = "Date not provided"
        gig_venue = "Venue not provided"
        band_name = "Band not provided"

    # Maximum duration for each set in seconds (45 minutes = 2700 seconds)
    max_set_duration = 45 * 60
    max_songs_per_page = 10  # Estimate how many songs can fit on a page at default font size
    min_font_size = 10  # Minimum font size allowed for shrinking
    default_font_size = 18

    total_songs = len(df)
    current_set_time = 0
    set_songs = []
    total_time = 0
    set = 1
    
    for index, row in df.iterrows():
        song = row['Song']
        length = row['Length']
        song_duration = time_to_seconds(length)

        # Check if adding the current song would exceed the set time limit
        if current_set_time + song_duration > max_set_duration:
            # Write the current set to the document and adjust font size if needed
            # Add the gig details at the start of the document            
            add_gig_details(doc, band_name, gig_venue, gig_date)
            write_set_to_document(set, doc, set_songs, max_songs_per_page, default_font_size, min_font_size)
            doc.add_page_break()
            set=set+1

            # Start a new set
            set_songs = []
            current_set_time = 0

        # Add the song to the current set
        set_songs.append({'Song': song, 'Length': length})
        current_set_time += song_duration
        total_time += song_duration

    # Write the remaining songs to the document (last set)
    if set_songs:
        # Add the gig details at the start of the document
        set+=set        
        add_gig_details(doc, band_name, gig_venue, gig_date)
        write_set_to_document(set, doc, set_songs, max_songs_per_page, default_font_size, min_font_size)

    # Save the document
    doc.save('setlist.docx')
    print("Setlist generated and saved as 'setlist.docx'")

# Function to write a set to the document, adjusting the font size if necessary
def write_set_to_document(set, doc, set_songs, max_songs_per_page, default_font_size, min_font_size):
    num_songs = len(set_songs)
    font_size = default_font_size
    
    add_set_number(doc, set)

    # If the set has more songs than can fit on a page, reduce the font size
    if num_songs > max_songs_per_page:
        font_size = max(min_font_size, default_font_size - (num_songs - max_songs_per_page))
    logging.debug(f"Font Size: {font_size}")

    total_time_in_set = 0
    song_number = 1
    for song in set_songs:        
        add_song_to_document(doc, song_number, song['Song'], song['Length'], font_size)
        total_time_in_set += time_to_seconds(song['Length'])
        song_number = song_number+1

    # Add the total set length at the bottom of the page
    total_minutes, total_seconds = divmod(total_time_in_set, 60)
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run(f"Total Set Length: {total_minutes} min {total_seconds} sec")
    run.font.size = Pt(font_size - 4)  # Slightly smaller font for the set length summary

# Main function to handle arguments and execution
if __name__ == "__main__":
    parser = argparse.ArgumentParser(description='Generate a setlist from a CSV of songs.')
    parser.add_argument('csv_file', type=str, help='Path to the CSV file containing song names and lengths.')
    parser.add_argument('--config_file', type=str, help='Optional config file with gig details (YAML format)', required=False)

    args = parser.parse_args()

    logging.basicConfig(level=logging.DEBUG)    
    
    # Generate the setlist with optional config file
    create_setlist(args.csv_file, args.config_file)
